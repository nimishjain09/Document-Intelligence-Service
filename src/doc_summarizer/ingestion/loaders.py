"""Concrete document loaders for supported file types."""

from __future__ import annotations

from docx import Document as DocxDocument
from pypdf import PdfReader

from doc_summarizer.config.logging_config import get_logger
from doc_summarizer.core.exceptions import DocumentLoadError
from doc_summarizer.ingestion.base import DocumentLoader
from doc_summarizer.ingestion.models import LoadedDocument

logger = get_logger(__name__)


class TxtLoader(DocumentLoader):
    """Loader for plain-text (.txt) files."""

    def load(self) -> LoadedDocument:
        try:
            text = self.path.read_text(encoding="utf-8", errors="ignore")
            logger.info("Loaded TXT: %s", self.path.name)
            return LoadedDocument(source_path=self.path, text=text)
        except OSError as exc:
            raise DocumentLoadError(f"Failed to read TXT '{self.path}': {exc}") from exc


class PdfLoader(DocumentLoader):
    """Loader for PDF (.pdf) files."""

    def load(self) -> LoadedDocument:
        try:
            reader = PdfReader(str(self.path))
            pages = [page.extract_text() or "" for page in reader.pages]
            text = "\n".join(pages)
            logger.info("Loaded PDF: %s (%d pages)", self.path.name, len(reader.pages))
            return LoadedDocument(source_path=self.path, text=text)
        except Exception as exc:
            raise DocumentLoadError(f"Failed to read PDF '{self.path}': {exc}") from exc


class DocxLoader(DocumentLoader):
    """Loader for Word (.docx) files."""

    def load(self) -> LoadedDocument:
        try:
            document = DocxDocument(str(self.path))
            text = "\n".join(p.text for p in document.paragraphs)
            logger.info("Loaded DOCX: %s", self.path.name)
            return LoadedDocument(source_path=self.path, text=text)
        except Exception as exc:
            raise DocumentLoadError(f"Failed to read DOCX '{self.path}': {exc}") from exc